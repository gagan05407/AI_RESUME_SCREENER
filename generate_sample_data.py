import os
import docx

def create_sample_docx(filepath, title, contact, summary, experience_lines, education_lines, skills_line):
    """
    Programmatically creates a formatted DOCX resume.
    """
    doc = docx.Document()
    
    # Title / Header
    doc.add_heading(title, level=0)
    doc.add_paragraph(contact)
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)
    
    # Experience
    doc.add_heading("Professional Experience", level=1)
    for line in experience_lines:
        doc.add_paragraph(line)
        
    # Education
    doc.add_heading("Education", level=1)
    for line in education_lines:
        doc.add_paragraph(line)
        
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(skills_line)
    
    # Save
    doc.save(filepath)
    print(f"Created sample resume: {filepath}")

def main():
    # Make directories
    os.makedirs("data/resumes", exist_ok=True)
    os.makedirs("data/job_descriptions", exist_ok=True)
    
    # 1. John Doe - Machine Learning Engineer
    create_sample_docx(
        filepath="data/resumes/john_doe_ml_engineer.docx",
        title="John Doe",
        contact="Email: john.doe@email.com | Phone: +1-555-0199 | GitHub: github.com/johndoe-ml",
        summary="Passionate Machine Learning Engineer with 4 years of experience building and deploying scalable deep learning models. Expert in NLP, computer vision, and building microservices around predictive models.",
        experience_lines=[
            "Senior ML Engineer | TechCorp Inc. (2022 - Present)",
            "- Designed and implemented an NLP translation service using PyTorch and Transformer architectures.",
            "- Containerized applications using Docker and deployed to AWS ECS, saving 30% in infrastructure costs.",
            "- Engineered automated data processing pipelines with Python and SQL, handling 10M+ daily events.",
            "ML Engineer | DataSoft Solutions (2020 - 2022)",
            "- Built classification and regression models using Scikit-Learn and Pandas.",
            "- Maintained SQL databases, optimizing index speeds and querying routines."
        ],
        education_lines=[
            "M.Tech in Computer Science | Technical University of India (2018 - 2020)",
            "B.Tech in Information Technology | Mumbai University (2014 - 2018)"
        ],
        skills_line="Python, Machine Learning, Deep Learning, PyTorch, TensorFlow, Scikit-Learn, Pandas, NumPy, SQL, Docker, AWS, Git, NLP"
    )
    
    # 2. Jane Smith - Full Stack Developer
    create_sample_docx(
        filepath="data/resumes/jane_smith_fullstack.docx",
        title="Jane Smith",
        contact="Email: jane.smith@email.com | Phone: +1-555-0144 | Web: janesmith.dev",
        summary="Creative Full Stack Developer with 2 years of experience specializing in high-performance web applications. Extensive experience in JavaScript/TypeScript ecosystems, building responsive user interfaces, and designing RESTful APIs.",
        experience_lines=[
            "Full Stack Software Engineer | WebFlow Studio (2024 - Present)",
            "- Architected client-side features in React and TypeScript, improving web accessibility and mobile responsiveness.",
            "- Constructed scalable server-side systems with Node.js, Express, and MongoDB.",
            "- Created clean database schemas in PostgreSQL and MySQL to manage user profiles.",
            "Junior Web Developer | Innovate Web Corp (2022 - 2024)",
            "- Developed layouts using HTML and CSS.",
            "- Configured automated testing and version control flows with Git."
        ],
        education_lines=[
            "B.Tech in Computer Engineering | Delhi Technological University (2018 - 2022)"
        ],
        skills_line="JavaScript, TypeScript, React, Node.js, Express, HTML, CSS, Git, PostgreSQL, MongoDB, SQL, Flask"
    )
    
    # 3. Alex Lee - Junior Data Analyst
    create_sample_docx(
        filepath="data/resumes/alex_lee_data_analyst.docx",
        title="Alex Lee",
        contact="Email: alex.lee@email.com | Phone: +1-555-0122",
        summary="Detail-oriented Data Analyst with 1 year of experience transforming complex datasets into actionable business intelligence. Fluent in statistical computing, data cleaning, and creating dashboard tools.",
        experience_lines=[
            "Data Analyst | RetailMetrics (2025 - Present)",
            "- Automated weekly business reporting using Pandas and NumPy, reducing manual reporting workload by 50%.",
            "- Formulated complex SQL query logic to pull monthly sales trends from databases containing millions of transactions.",
            "- Crafted interactive dashboards to display customer retention metrics."
        ],
        education_lines=[
            "Bachelor of Science (B.Sc) in Statistics | Pune University (2021 - 2024)"
        ],
        skills_line="Python, SQL, Pandas, NumPy, Scikit-Learn, Git, HTML, CSS"
    )
    
    # 4. Job Description - ML Engineer
    jd_content = """Job Title: Senior Machine Learning Engineer

About the Role:
We are looking for a Senior Machine Learning Engineer to join our Core AI team. You will lead the design, implementation, and deployment of semantic search and NLP models.

Required Skills & Qualifications:
- 3+ years of professional experience in machine learning or software engineering.
- Deep expertise in Python, SQL, and Git.
- Practical experience with Deep Learning frameworks, specifically PyTorch or TensorFlow.
- Experience with containerization technologies like Docker.
- Strong educational background: Bachelor's (B.Tech/B.E./B.S.) or Master's (M.Tech/M.S.) degree in Computer Science or related fields.

Preferred Skills & Qualifications:
- Familiarity with cloud platforms, especially AWS.
- Experience deploying machine learning applications using Streamlit or Flask.
- Knowledge of NLP, BERT, and Large Language Models (LLMs).

Join us in building the future of recruitment tech!
"""
    
    jd_path = "data/job_descriptions/ml_engineer_jd.txt"
    with open(jd_path, "w", encoding="utf-8") as f:
        f.write(jd_content)
    print(f"Created sample Job Description: {jd_path}")

if __name__ == "__main__":
    main()
