import os
import io
import re
from typing import Union, BinaryIO
import pdfplumber
import docx

def extract_text_from_pdf(file_source: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts raw text from a PDF file.
    
    Args:
        file_source: Absolute path to the PDF, a bytes object, or a file-like stream.
        
    Returns:
        Extracted raw text from the PDF.
    """
    text = ""
    try:
        if isinstance(file_source, bytes):
            file_source = io.BytesIO(file_source)
            
        with pdfplumber.open(file_source) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file_source: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts raw text from a DOCX file.
    
    Args:
        file_source: Absolute path to the DOCX, a bytes object, or a file-like stream.
        
    Returns:
        Extracted raw text from the DOCX.
    """
    text = ""
    try:
        if isinstance(file_source, bytes):
            file_source = io.BytesIO(file_source)
            
        doc = docx.Document(file_source)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
                
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text += " | ".join(filter(None, row_text)) + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def clean_text(text: str) -> str:
    """
    Cleans raw text by normalizing whitespaces, removing non-printable characters,
    and removing excessive spacing.
    
    Args:
        text: Raw text.
        
    Returns:
        Clean, normalized text.
    """
    if not text:
        return ""
    # Normalize whitespaces to single spaces
    text = re.sub(r'\s+', ' ', text)
    # Strip leading/trailing whitespaces
    return text.strip()

def parse_resume(file_source: Union[str, bytes, BinaryIO], file_name: str = "") -> str:
    """
    Parses a resume (PDF/DOCX) and returns the raw extracted text.
    
    Args:
        file_source: File path, bytes, or file-like object.
        file_name: Name of the file (required if file_source is bytes/file-like to infer extension).
        
    Returns:
        Extracted raw text.
    """
    if not file_name and isinstance(file_source, str):
        file_name = file_source

    file_name_lower = file_name.lower()
    
    if file_name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_source)
    elif file_name_lower.endswith(".docx"):
        return extract_text_from_docx(file_source)
    else:
        # Fallback: try to read as plain text
        try:
            if isinstance(file_source, str):
                with open(file_source, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            elif isinstance(file_source, bytes):
                return file_source.decode("utf-8", errors="ignore")
            else:
                return file_source.read().decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"Unsupported file format and fallback failed: {e}")
            return ""
